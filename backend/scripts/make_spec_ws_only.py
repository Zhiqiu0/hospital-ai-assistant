# -*- coding: utf-8 -*-
"""把 HIS 接口规范收敛为单一 WebSocket 通道（v1.3 双通道 → v1.4 单通道）。

背景（2026-08-17，用户拍板「母本也删，不要双通道了」）：
  原规范是「方案 A（HTTPS 直连）+ 方案 B（WebSocket 长连接）二选一」的双通道设计，
  厂商早已确认走方案 B。方案 A 的残留散在正文/表格/附录十余处，之前是「保留+加注
  适用范围」，对厂商全是「读半天发现跟我无关」的噪音。本脚本把方案 A 彻底删除、
  把「双通道二选一」的框架改写成「单一 WebSocket 通道」，产出 v1.4。

做法：读当前 v1.3 docx（业务字段表等主体内容已定稿，全部保留不动），只改传输框架
  相关的段落/表格行 + 删方案 A 专属内容 + 修 40007 错误码行掉字体的 bug。产物是新的
  唯一真源；双通道的 v1.2 母本 + v1.3 生成管线就此退役为历史。

用法：python scripts/make_spec_ws_only.py
"""
from pathlib import Path

import docx
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import Table

SRC = Path("d:/Code/hospital-ai-assistant/docs/MediScribe_HIS接口规范_v1.3.docx")
OUT = Path("d:/Code/hospital-ai-assistant/docs/MediScribe_HIS接口规范_v1.4_单通道.docx")
FONT = "Microsoft YaHei"

applied: list[str] = []
failed: list[str] = []


def _set_font(run, size=9.0):
    """统一字体口径（与错误码表其余行一致：Microsoft YaHei 9.0），修 40007 掉字体。"""
    run.font.size = Pt(size)
    run.font.name = FONT
    rpr = run._r.get_or_add_rPr()
    fonts = rpr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(fonts)
    for a in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        fonts.set(qn(a), FONT)


def _find(doc, needle):
    for p in doc.paragraphs:
        if needle in p.text:
            return p
    return None


def _rewrite(doc, needle, new_text, tag):
    """整段替换为单通道措辞，保留首个 run 的字体，不加变更高亮（这是新基线不是待审改动）。"""
    p = _find(doc, needle)
    if p is None:
        failed.append(f"{tag}｜找不到段落：{needle[:20]}")
        return
    runs = list(p.runs)
    if not runs:
        p.add_run(new_text)
    else:
        runs[0].text = new_text
        for r in runs[1:]:
            r._r.getparent().remove(r._r)
    applied.append(tag)


def _sub(doc, anchor, old, new, tag):
    """段内子串替换（跨 run 合并后替换），保留段落其余内容。"""
    p = _find(doc, anchor)
    if p is None:
        failed.append(f"{tag}｜找不到段落：{anchor[:20]}")
        return
    full = "".join(r.text for r in p.runs)
    if old not in full:
        failed.append(f"{tag}｜段内找不到子串：{old[:20]}")
        return
    runs = list(p.runs)
    runs[0].text = full.replace(old, new)
    for r in runs[1:]:
        r._r.getparent().remove(r._r)
    applied.append(tag)


def _strip_http(doc, needle, tag):
    """去掉附录 A 示例里的「POST url / Headers: X-…」外壳，只留从第一个 { 起的业务 JSON。

    单通道下这些接口都走 WebSocket 消息 payload（外层信封见 A.4 / 7.3），HTTP 请求头
    是方案 A 残留。保留 JSON 正文（业务字段是通道无关的），前面补一行归属说明。
    """
    p = _find(doc, needle)
    if p is None:
        failed.append(f"{tag}｜找不到示例段：{needle[:20]}")
        return
    full = "".join(r.text for r in p.runs)
    lines = full.split("\n")
    kept, started = [], False
    for ln in lines:
        if not started and ln.strip().startswith("{"):
            started = True
        if started:
            kept.append(ln)
    body_json = "\n".join(kept) if kept else full
    new_text = "▲ （作为 WebSocket 消息的 payload 发送，外层信封见 A.4 / 7.3）\n" + body_json
    runs = list(p.runs)
    runs[0].text = new_text
    for r in runs[1:]:
        r._r.getparent().remove(r._r)
    applied.append(tag)


def _delete(doc, needle, tag):
    p = _find(doc, needle)
    if p is None:
        failed.append(f"{tag}｜找不到待删段落：{needle[:20]}")
        return
    p._p.getparent().remove(p._p)
    applied.append(tag)


def _tables(doc):
    out = []
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:tbl"):
            out.append(Table(child, doc))
    return out


def _cell_set(tbl, row_key, col, text, tag, *, key_col=0, size=9.0):
    for row in tbl.rows:
        if row.cells[key_col].text.strip().replace("▲ ", "").startswith(row_key):
            cell = row.cells[col]
            para = cell.paragraphs[0]
            for r in list(para.runs):
                r._r.getparent().remove(r._r)
            run = para.add_run(text)
            _set_font(run, size)
            applied.append(tag)
            return
    failed.append(f"{tag}｜表里找不到行：{row_key}")


def _row_delete(tbl, row_key, tag, *, key_col=0):
    for row in tbl.rows:
        if row.cells[key_col].text.strip().replace("▲ ", "").startswith(row_key):
            row._tr.getparent().remove(row._tr)
            applied.append(tag)
            return
    failed.append(f"{tag}｜表里找不到待删行：{row_key}")


def main() -> int:
    doc = docx.Document(str(SRC))

    # ── 版本号 v1.3 → v1.4（标题/页脚/摘要处）────────────────────────────────
    for p in doc.paragraphs:
        if "v1.3" in p.text and ("版本" in p.text or "接口规范" in p.text or "评审" in p.text):
            for r in p.runs:
                r.text = r.text.replace("v1.3", "v1.4")

    # ── 正文单通道改写 ────────────────────────────────────────────────────────
    R = [
        ("传输通道已双方确认采用方案 B",
         "说明：本对接为医院授权的点对点直连，不接入省级平台；消息一律为 JSON。"
         "传输通道为 WebSocket 长连接（wss，由贵方 HIS 侧程序主动向我方建立，见第 7 章），"
         "接诊推送与病历回写复用同一条连接双向收发。不依赖动态库（dll）。", "1.3 通道说明"),
        ("HIS 将接诊信息(就诊流水号+患者+医生工号) POST 到 MediScribe 公网",
         "▲ ①医生在 HIS 接诊病人\n"
         "      │  HIS 将接诊信息(就诊流水号+患者+医生工号)经 wss 长连接推送到 MediScribe\n"
         "      ▼\n"
         "②MediScribe 自动建档并按 doctor_code 派入该医生的浏览器工作台；医生点击叫号提示进入\n"
         "      │  医生借 AI 完成病历并确认\n"
         "      ▼\n"
         "③MediScribe 沿同一 wss 连接下发病历回写(写入) → 再下发触发刷新\n"
         "      │  以同一就诊流水号关联\n"
         "      ▼\n"
         "④HIS 病历页面刷新，显示待确认病历 → 医生核对保存", "1.3 流程图"),
        ("方案 A 下两个方向均采用下述签名鉴权",
         "注：此凭证为本点对点对接专用，与省平台向贵方下发的凭证无关，无需经省平台。"
         "凭证仅需我方分配给贵方一套 appId/appSecret，接诊推送与病历回写共用（WebSocket "
         "单连接双向复用，见 7.2）。如需加强，可在签名之上叠加内网 IP 白名单。", "2.2 凭证"),
        ("双方约定本对接专用凭证 appId + appSecret，由「接收方」为「调用方」分配",
         "双方约定本对接专用凭证 appId + appSecret，由 MediScribe 统一分配给贵方一套；"
         "接诊推送与病历回写共用（WebSocket 单连接双向复用，见 7.2）。", "2.2 凭证分配"),
        ("每个请求在 HTTP Header 携带以下公共参数",
         "鉴权公共参数（在 WebSocket 握手 URL query 与消息信封中携带，承载位置与待签串"
         "组成见第 7 章 7.2 / 7.3）：", "2.2 参数承载"),
        ("待签串 = X-App-Id + X-Timestamp + X-Nonce + 请求body原文",
         "待签串 = app_id + timestamp + nonce + 报文原文；"
         "X-Sign = Hex( HMAC_SHA256( appSecret, 待签串 ) )。"
         "握手与业务消息的报文原文取值分别见 7.2、7.3（按实际发出的字节验签）。", "2.2 待签串"),
        ("方案 B 补充：WebSocket 通道的业务消息不做 nonce 去重",
         "防重放：WebSocket 通道的业务消息不做 nonce 去重，依赖握手鉴权 + ±5 分钟时间戳窗口 "
         "+ msg_id 幂等；重发时 timestamp/nonce/sign 重新生成而 msg_id 保持不变（见 7.4）。"
         "另请双方保证服务器 NTP 校时，避免时间戳超出 ±5 分钟窗口。", "2.2 防重放"),
        ("贵方 HIS 直接 POST 到 MediScribe 公网地址即可",
         "接诊推送沿贵方主动建立的 wss 长连接上行（见第 7 章），贵方无需对外开放任何接口；"
         "仅需贵方 HIS 侧程序能主动出网访问 wss://mediscribe.cn（443 端口出向连接）。", "5.1 网络"),
        ("回写的写入与刷新接口由贵方提供，按其网络可达性两种方式",
         "病历回写与触发刷新沿同一条 wss 连接下行，贵方无需提供公网回写地址、无需对外开放"
         "任何接口，我方也无需连入院内网、不在院内部署任何程序。", "5.2 网络"),
        ("5.3 方案 B（WebSocket）下的网络简化", "5.3 网络要求小结", "5.3 标题"),
        ("若采用 WebSocket 通道，5.1 与 5.2 的网络要求整体简化",
         "综上：贵方 HIS 侧程序只需能主动出网访问 wss://mediscribe.cn（443 端口出向连接）即可，"
         "无需对外开放任何接口、无需提供公网回写地址；我方亦无需在内网中转，院内不部署我方任何程序。",
         "5.3 小结"),
        ("凭证交换：按所选通道分配 appId/appSecret",
         "1. 凭证交换：由我方分配一套 appId/appSecret 给贵方（接诊推送与回写共用），"
         "签名统一 HMAC-SHA256。", "6 凭证交换"),
        ("3. 通道连通（方案 B）",
         "▲ 3. 通道连通：贵方客户端 wss 握手鉴权通过；上行 encounter_admit 我方回 ack code=0 "
         "且患者出现在医生工作台队列；下行 record_writeback / record_refresh 贵方落库并回 ack；"
         "ping/pong 心跳与断线自动重连正常。", "6 通道连通"),
        ("7. WebSocket 长连接通道（方案 B）", "7. WebSocket 长连接通道", "7 标题"),
        ("由贵方 HIS 侧程序（院内）主动向我方建立一条加密长连接",
         "由贵方 HIS 侧程序（院内）主动向我方建立一条加密长连接（wss）：接诊推送沿此连接上行，"
         "病历回写与刷新沿同一连接下行。业务字段、必填性、幂等与错误码约定与第 2、3 章完全一致，"
         "仅传输管道不同。贵方无需对外开放任何接口，我方也无需连入院内网。第 3 章三个接口均经"
         "本通道传输。", "7.1 说明"),
        ("A.4 WebSocket 消息示例（方案 B）", "A.4 WebSocket 消息示例", "A.4 标题"),
        ("【已确认】传输通道：方案 B（WebSocket 长连接）",
         "【已确认】传输通道：WebSocket 长连接（wss，由贵方 HIS 侧主动建立，见第 7 章）。", "附B 通道"),
        ("【已确认】凭证：方案 B 下仅我方分配一套",
         "【已确认】凭证：仅我方分配一套 appId/appSecret（已生成，线下交付贵方）；贵方无需为我方分配。",
         "附B 凭证"),
        ("— 本文档 v1.4，供双方评审。范围：接诊数据推送 + 病历回写",
         "— 本文档 v1.4，供双方评审。范围：接诊数据推送 + 病历回写（传输通道：WebSocket 长连接）。 —",
         "文末范围"),
    ]
    for needle, new, tag in R:
        _rewrite(doc, needle, new, tag)

    # ── 删除方案 A 专属段落 ───────────────────────────────────────────────────
    _delete(doc, "通道说明：上方流程图按方案 A", "删·流程图双通道注")
    _delete(doc, "重试规则（本条适用于方案 A）", "删·方案A重试规则")
    _delete(doc, "（本节原「诊室部署极薄 Agent 做内网中转」方案已取消）方案 B 下我方不调用", "删·回落方案A回写地址")

    # ── 表格改动 ──────────────────────────────────────────────────────────────
    tbls = _tables(doc)
    # T1 传输表：协议/方法/通道 改为 WebSocket 口径
    _cell_set(tbls[1], "协议", 1, "WebSocket over TLS（wss）", "T1 协议")
    _cell_set(tbls[1], "方法", 1, "消息帧：Text 帧（UTF-8 文本，承载 JSON）", "T1 方法")
    _cell_set(tbls[1], "通道", 1, "WebSocket 长连接（wss，由贵方 HIS 侧主动建立，见第 7 章）", "T1 通道")
    # T4 错误码：删 40006，修 40007 掉字体
    _row_delete(tbls[4], "40006", "删·40006")
    _cell_set(tbls[4], "40007", 0, "40007", "T4 40007 code")
    _cell_set(tbls[4], "40007", 1,
              "业务处理失败（具体原因见 message，如接诊推送的医生工号未在 MediScribe 注册）",
              "T4 40007 含义")
    # T12 回写可达性表整表已无意义（单通道回写沿 wss 下行）——删「仅院内网可达」行，
    # 保留表头+公网行改写为 WS 口径
    _row_delete(tbls[12], "仅院内网可达", "删·T12 院内网行")
    _cell_set(tbls[12], "公网可达", 0, "WebSocket 下行", "T12 可达性")
    _cell_set(tbls[12], "WebSocket 下行", 1,
              "回写与刷新沿贵方建立的同一条 wss 连接下行，我方按连接自动路由回来源诊室，"
              "贵方无需提供任何公网地址", "T12 实现")

    # ── 第二轮：签名参数表 / 附录 A HTTP 示例 / 摘要 / 零散「方案 B」字样 ──────────
    # 修订摘要（正文第一屏，厂商最先读）：改为以「通道收敛」为头条
    _rewrite(doc, "v1.3 修订：共 32 处",
             "▲ v1.4 修订：传输通道收敛为单一 WebSocket 长连接，删除原方案 A（HTTPS 直连）"
             "相关内容——原方案 A 仅为备选，双方已确认采用 WebSocket，本版不再并列双通道，"
             "以免造成「读了才发现与己无关」的干扰。业务字段、必填性、幂等与错误码约定均未改变。"
             "其余需请您过目的要点见文末《附录 C》。正文中修订过的位置以句首「▲」+ 淡底色标出。",
             "P4 修订摘要")
    # T2 签名参数表：X- HTTP 头名 → 与 7.2/7.3 一致的 app_id/timestamp/nonce/sign
    _cell_set(tbls[2], "Header", 0, "参数", "T2 表头")
    _cell_set(tbls[2], "X-App-Id", 0, "app_id", "T2 app_id")
    _cell_set(tbls[2], "X-Timestamp", 0, "timestamp", "T2 timestamp")
    _cell_set(tbls[2], "X-Nonce", 0, "nonce", "T2 nonce")
    _cell_set(tbls[2], "X-Sign", 0, "sign", "T2 sign")
    # 2.2 待签串里残留的 X-Sign 写法
    _rewrite(doc, "待签串 = app_id + timestamp + nonce + 报文原文；X-Sign",
             "待签串 = app_id + timestamp + nonce + 报文原文；"
             "sign = Hex( HMAC_SHA256( appSecret, 待签串 ) )。"
             "握手与业务消息的报文原文取值分别见 7.2、7.3（按实际发出的字节验签）。", "2.2 待签串sign")
    # 7.3 信封表引用的 X-App-Id
    _cell_set(tbls[13], "app_id", 2, "同 2.2 的 app_id", "T13 app_id引用", key_col=0)
    # 零散「方案 B」字样
    _rewrite(doc, "凭证由我方分配（方案 B 为单向一套",
             "▲ 请贵方提供：测试环境的 HIS 客户端与测试用就诊数据；凭证由我方分配（单向一套，"
             "贵方无需为我方分配）。我方为单环境部署（即 wss://mediscribe.cn），联调与正式使用"
             "同一地址、同一套凭证，无需另配测试地址与测试凭证。因此请贵方联调时务必使用"
             "「虚构的」就诊数据（已签发病历产品上不可作废，会永久留在生产库）。", "P69 方案B字样")
    _cell_set(tbls[6], "agent_device_ip", 3,
              "▲ 选填，仅作排障留痕。回写按 WebSocket 连接自动路由回来源诊室，不依赖本字段，可不填",
              "T6 agent_ip", key_col=0)
    # 附录 A.1~A.3：去掉 HTTP POST + X-头 外壳，只留业务 JSON（信封见 A.4 / 7.3）
    _strip_http(doc, "POST https://mediscribe.cn/api/v1/his/encounter/admit", "A.1 去HTTP壳")
    _strip_http(doc, "POST http://{his_host}:{port}/{writeback_path}", "A.2 去HTTP壳")
    _strip_http(doc, "POST http://{his_host}:{port}/{refresh_path}", "A.3 去HTTP壳")
    # 附录 C 标题
    _rewrite(doc, "附录 C：本次修订说明（v1.2 → v1.3）",
             "附录 C：本次修订说明（v1.2 → v1.4）", "附C 标题")

    # 7.2 里残留的「方案 B 下仅需一套凭证」（单通道无需再点出方案 B）
    _sub(doc, "消息由我方自动路由回该就诊接诊推送的来源连接",
         "凭证简化：方案 B 下仅需一套凭证", "凭证：仅需一套凭证", "7.2 凭证简化")

    # ── 附录 C 修订说明：清掉单元格里的「方案 A/B」措辞（内容保留）──────────────
    # 顺序：先长后短，先修带空格/上下文的，再兜底泛化，避免替换后留「方案 的」这种碎片
    REPL = [
        ("方案 A（诊室 Agent 中转）时代的设计", "早期诊室 Agent 中转方案的设计（已废弃）"),
        ("方案 A 时代的方案", "早期方案（已废弃）"),
        ("方案 A 时期", "早期（诊室 Agent 方案）时期"),
        ("方案 A 的", "早期 HTTP 直连方案的"),
        ("方案 A", "早期 HTTP 直连方案"),
        ("通道定为方案 B 后", "通道定为 WebSocket 后"),
        ("方案 B 下", "WebSocket 通道下"),
        ("方案 B", "WebSocket 通道"),
    ]
    for tbl in (tbls[15], tbls[16]):
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for r in para.runs:
                        if "方案" in r.text:
                            for old, new in REPL:
                                r.text = r.text.replace(old, new)

    doc.save(str(OUT))
    print(f"✅ 已生成：{OUT.name}")
    print(f"应用 {len(applied)} 处，失败 {len(failed)} 处")
    for f in failed:
        print("  ✗ " + f)
    return 1 if failed else 0


if __name__ == "__main__":
    raise SystemExit(main())
