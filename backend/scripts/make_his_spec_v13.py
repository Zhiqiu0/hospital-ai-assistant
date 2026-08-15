"""生成 v1.3 变更版规范：开头插一页变更摘要。

两个要点：

1. **字体必须三个属性一起设**（ascii / hAnsi / eastAsia）。
   python-docx 的 `run.font.name` 只设 ascii 与 hAnsi，中文会掉回主题默认字体，
   于是中英文各用一套字体，排版看着很乱。原文档正文用的是 Microsoft YaHei
   且三属性齐全，这里保持一致。

2. **不标注「贵方是否需改代码」**。那样写像是在给对方派活；
   末列改为「背景说明」，交代这条从何而来，厂商自己判断要不要动代码。
   注意分寸（2026-08-15 调整）：如实说明 ≠ 逐条自我检讨。
   原先列名叫「为何此前未写明」、正文还散着 11 处「属我方疏漏」，
   30 行叠起来就成了一份检讨书。事实照写，自我定性删掉；
   开场白也从「我们发现不一致」改成「赶在贵方动手之前一次性定清楚」——
   厂商怕的从来不是文档改了多少处，而是写完了你又来改。
"""
import copy
import sys
from pathlib import Path

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

sys.path.insert(0, str(Path(__file__).parent))
from his_spec_changes_v13 import CHANGES  # noqa: E402

SRC = Path("d:/Code/hospital-ai-assistant/docs/MediScribe_HIS接口规范.docx")
OUT = Path("d:/Code/hospital-ai-assistant/docs/MediScribe_HIS接口规范_v1.3.docx")

FONT = "Microsoft YaHei"          # 与原文档正文一致
RED = RGBColor(0xC0, 0x00, 0x00)
DARK = RGBColor(0x33, 0x33, 0x33)
GRAY = RGBColor(0x66, 0x66, 0x66)

PRIO = {
    "A": ("A 类｜请在贵方开发完成前对齐", RGBColor(0xC0, 0x00, 0x00)),
    "B": ("B 类｜原要求已取消，贵方可少做", RGBColor(0x00, 0x70, 0xC0)),
    "C": ("C 类｜口径补充，避免联调时产生分歧", RGBColor(0x44, 0x60, 0x70)),
}


def style_run(run, *, size=10.5, bold=False, color=DARK):
    """统一设字体——ascii / hAnsi / eastAsia 三个都要设，否则中文会掉默认字体。"""
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(fonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        fonts.set(qn(attr), FONT)
    return run


def para_before(anchor, text="", *, size=10.5, bold=False, color=DARK,
                align=None, space_after=4, line=None):
    new_p = copy.deepcopy(anchor._p)
    anchor._p.addprevious(new_p)
    from docx.text.paragraph import Paragraph
    p = Paragraph(new_p, anchor._parent)
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    p.text = ""
    style_run(p.add_run(text), size=size, bold=bold, color=color)
    p.paragraph_format.space_after = Pt(space_after)
    if line:
        p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    return p


def fill_cell(cell, text, *, bold=False, color=DARK, size=9):
    para = cell.paragraphs[0]
    for r in list(para.runs):
        r._r.getparent().remove(r._r)
    style_run(para.add_run(text), size=size, bold=bold, color=color)
    para.paragraph_format.space_after = Pt(2)


def main() -> None:
    doc = docx.Document(str(SRC))

    for p in doc.paragraphs[:5]:
        if "版本 v1.2" in p.text:
            for r in p.runs:
                r.text = r.text.replace("v1.2", "v1.3").replace("2026-08-11", "2026-08-14")

    anchor = doc.paragraphs[0]

    para_before(anchor, "本次修订说明（v1.2 → v1.3）", size=15, bold=True,
                color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    # 定调（2026-08-15 改）：同一份清单，说法不同性质完全不同。
    # 原文是「我们发现规范与实现不一致、大多是我方没同步」——一上来就道歉，
    # 30 行看下来像一份检讨。厂商真正怕的从来不是「文档改了 30 处」，
    # 而是「我照着写完了你又来改」；而我们恰好赶在他动手之前。
    # 所以开场白讲的应该是这件事的价值：都在你写第一行代码之前定下来了。
    para_before(
        anchor,
        "老师好。在贵方动手开发之前，我方先把系统代码与本规范做了一轮逐条核对，"
        "把所有会影响贵方实现的地方一次性梳理清楚，共 %d 处，全部列在下表。"
        "这样做只为一件事：把该定的都定在贵方写第一行代码之前，"
        "免得开发到一半再来回改。" % len(CHANGES),
        size=10.5, color=DARK, space_after=4, line=1.5,
    )
    para_before(
        anchor,
        "下表逐条列明：改了什么、为什么这样改、以及相关背景。"
        "其中 A 类 %d 项建议优先过目——这几处若两边理解不一致，"
        "联调时会表现为「消息发出去了但对方没反应」这类不易定位的问题。"
        "其余各项多为口径写明与要求取消，不涉及额外开发量。"
        % len([c for c in CHANGES if c["prio"] == "A"]),
        size=10.5, color=DARK, space_after=12, line=1.5,
    )

    for prio in ("A", "B", "C"):
        items = [c for c in CHANGES if c["prio"] == prio]
        if not items:
            continue
        title, color = PRIO[prio]
        para_before(anchor, f"{title}（{len(items)} 项）",
                    size=12, bold=True, color=color, space_after=6)

        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Table Grid"
        # 末列叫「背景说明」而不是「为何此前未写明」：内容一字不减，但列名不该
        # 把 30 行全框成认错。这一列真正的作用是让每条对厂商「安全」——
        # 告诉他这不是你漏做了什么、这条本期不涉及别去开发。
        for i, h in enumerate(["位置", "改了什么", "为什么这样改", "背景说明"]):
            fill_cell(tbl.rows[0].cells[i], h, bold=True, color=DARK)
        for c in items:
            row = tbl.add_row().cells
            fill_cell(row[0], c["loc"], bold=True)
            fill_cell(row[1], c["what"])
            fill_cell(row[2], c["why"])
            fill_cell(row[3], c["origin"], color=GRAY)
        anchor._p.addprevious(tbl._tbl)
        para_before(anchor, "", size=6, space_after=8)

    # 必须说明 ▲ 是什么：全文 31 处三角标记，不解释的话厂商会当成排版残留或乱码。
    # 另：实际用的是黄色高亮（w:highlight）而不是底纹（w:shd），措辞要对上。
    para_before(anchor,
                "以下为规范正文（v1.3）。正文中对应上述条目的位置已同步修订，"
                "均以黄色高亮 + 句首「▲」标出，便于对照。",
                size=10, color=GRAY, space_after=2)
    para_before(anchor, "─" * 52, size=9, color=GRAY, space_after=12)

    doc.save(str(OUT))
    print(f"已生成：{OUT}")
    for p in ("A", "B", "C"):
        print(f"  {p} 类 {len([c for c in CHANGES if c['prio']==p])} 项")


if __name__ == "__main__":
    main()
