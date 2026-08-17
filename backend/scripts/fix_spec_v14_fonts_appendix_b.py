# -*- coding: utf-8 -*-
"""v1.4 规范排版收口 + 附录 B 删两条（2026-08-18，用户看 Word 截图提出）。

排版（全是「同一张表里字体/字号不一致」这类客观缺陷，不动红字加粗等刻意强调）：
  ① 三张「接口信息」表（3.1/3.2/3.3）里没有 rPr 的 run —— Word 会按字符回落到不同字体，
     渲染成「经贵方建立的」几个字粗细不一的花脸（用户截图所指）。补 YaHei 9pt。
     顺带覆盖同表里同样没 rPr 的老行（接收处理）。
  ② 附录 C 标题段没有 rPr → 与附录 A/B 标题不同款。补成一样（YaHei 15pt 粗 深蓝）。
  ③ 表内字号混用：2.1 表「通道」行 10.5pt、3.1 字段表 vitals.*/allergy/past/dept_name
     10.5pt（其余行 9pt）→ 统一 9pt；7.3 信封表 3 个 ▲ 格 9pt（其余 10.5pt）→ 统一 10.5pt。
附录 B：
  ④ 删「▲ 补充：来源诊室连接失效时…」（7.2 正文与附录 C 已有，此处不是待确认项）
  ⑤ 删「▲【待确认】贵方客户端的开发语言…」（用户拍板：用什么语言厂商自己适配，不问）
  ⑥ 删字典项与 dept_code 项之间的空段
幂等：可重跑。
用法：python scripts/fix_spec_v14_fonts_appendix_b.py
"""
from pathlib import Path

import docx
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.table import Table

DOC = Path("d:/Code/hospital-ai-assistant/docs/MediScribe_HIS接口规范_v1.4_单通道.docx")
FONT = "Microsoft YaHei"
log: list[str] = []


def _set_font(run, size_pt: float, *, bold=None, color=None):
    """显式写 rFonts 四个槽 + 字号，杜绝按字符回落。"""
    run.font.name = FONT
    run.font.size = Pt(size_pt)
    rpr = run._r.get_or_add_rPr()
    fonts = rpr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, fonts)
    for a in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        fonts.set(qn(a), FONT)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def _sz(run):
    rpr = run._r.rPr
    if rpr is None:
        return None
    s = rpr.find(qn("w:sz"))
    return int(s.get(qn("w:val"))) if s is not None else None


def _tables(doc):
    return [Table(c, doc) for c in doc.element.body.iterchildren() if c.tag == qn("w:tbl")]


def _body_runs(tbl):
    for ri, row in enumerate(tbl.rows):
        if ri == 0:
            continue
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    if r.text.strip():
                        yield r


def main() -> int:
    doc = docx.Document(str(DOC))
    tables = _tables(doc)

    # ① 三张接口信息表 + ③ 表内字号统一 —— 规则：除 7.3/7.5 两张（10.5pt）外，表体一律 9pt
    for ti, tbl in enumerate(tables):
        is_ch7 = any(r.cells[0].text.strip() in ("type", "msg_id") for r in tbl.rows[1:2]) \
            or (tbl.rows[0].cells[0].text.strip() == "type")
        target = 21 if is_ch7 else 18   # 半磅
        n = 0
        for r in _body_runs(tbl):
            if r._r.rPr is None or _sz(r) != target:
                _set_font(r, target / 2)
                n += 1
        if n:
            log.append(f"表{ti}: 统一 {n} 个 run 到 {target/2}pt")

    # ② 附录 C 标题：字体同附录 A/B，且它原是居中的、A/B 是左对齐 → 一并对齐
    for p in doc.paragraphs:
        if p.text.strip().startswith("附录 C："):
            for r in p.runs:
                if r._r.rPr is None:
                    _set_font(r, 15, bold=True, color="1F4E79")
                    log.append("附录 C 标题补字体")
            if p.alignment is not None:
                p.alignment = None
                p.paragraph_format.space_after = None
                log.append("附录 C 标题改左对齐（同附录 A/B）")
            break

    # ④⑤⑥ 附录 B 删段
    kill_prefix = ("▲ 补充：来源诊室连接失效时", "▲ 【待确认】贵方客户端的开发语言")
    in_b = False
    for p in list(doc.paragraphs):
        t = p.text.strip()
        if t.startswith("附录 B："):
            in_b = True
            continue
        if t.startswith("附录 C："):
            break
        if not in_b:
            continue
        if t.startswith(kill_prefix) or t == "":
            p._p.getparent().remove(p._p)
            log.append(f"附录 B 删段: {t[:20] or '(空段)'}")

    doc.save(str(DOC))
    print("\n".join(log) if log else "（无改动，已是目标状态）")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
