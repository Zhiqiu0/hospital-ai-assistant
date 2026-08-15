"""把 23 项变更逐条应用到规范正文（v1.3）。

每处改动都：① 黄色高亮 ② 前缀「▲」标记，
让厂商翻到那一节时想漏都难（摘要页给全局，正文标记给细节）。

设计：所有修改都通过「按内容定位」而不是「按索引定位」——
索引会随插入变化，内容匹配失败会显式报错，不会静默改错地方。
"""
import re
import sys
from pathlib import Path

import docx
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.table import Table

sys.path.insert(0, str(Path(__file__).parent))
from his_spec_changes_v13 import CHANGES  # noqa: E402

FONT = "Microsoft YaHei"


def _set_font(run, *, size=None, bold=None, color=None):
    """三属性齐设（ascii/hAnsi/eastAsia），否则中文掉回默认字体、排版发乱。"""
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(fonts)
    for a in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        fonts.set(qn(a), FONT)
    return run

SRC = Path("d:/Code/hospital-ai-assistant/docs/MediScribe_HIS接口规范_v1.3.docx")
RED = RGBColor(0xC0, 0x00, 0x00)

_applied: list[str] = []
_failed: list[str] = []


# ── 基础操作 ────────────────────────────────────────────────────────────────

def _mark_runs(para, *, highlight=True):
    """给段落所有 run 打黄色高亮。"""
    for r in para.runs:
        if r.text.strip() and highlight:
            r.font.highlight_color = WD_COLOR_INDEX.YELLOW


def _set_text(para, text, *, mark=True):
    """整段替换文本，保留首个 run 的字体，并打标记。"""
    # 前缀幂等：同一段落可能被多条变更先后改到（如附录 A.1 示例要补 3 个字段），
    # 每次都无条件加前缀会叠成「▲ ▲ ▲ ▲」
    if mark and text.lstrip().startswith("▲"):
        mark_prefix = ""
    else:
        mark_prefix = "▲ " if mark else ""
    runs = list(para.runs)
    if not runs:
        r = para.add_run(text)
    else:
        runs[0].text = mark_prefix + text
        for r in runs[1:]:
            r._r.getparent().remove(r._r)
    _mark_runs(para, highlight=mark)


def _replace_in_para(para, old, new, *, tag):
    """段内子串替换（跨 run 时先合并再替换），命中才算成功。"""
    full = "".join(r.text for r in para.runs)
    if old not in full:
        _failed.append(tag)
        return False
    _set_text(para, full.replace(old, new))
    _applied.append(tag)
    return True


def _find_para(doc, needle):
    for p in doc.paragraphs:
        if needle in p.text:
            return p
    return None


def _edit_para(doc, needle, old, new, tag):
    p = _find_para(doc, needle)
    if p is None:
        _failed.append(f"{tag}（找不到段落）")
        return
    _replace_in_para(p, old, new, tag=tag)


def _rewrite_para(doc, needle, new_text, tag):
    p = _find_para(doc, needle)
    if p is None:
        _failed.append(f"{tag}（找不到段落）")
        return
    _set_text(p, new_text)
    _applied.append(tag)


def _append_para_after(doc, needle, text, tag):
    """在某段之后插入一段新的说明（红色加粗，作为补充条款）。"""
    import copy

    from docx.text.paragraph import Paragraph
    anchor = _find_para(doc, needle)
    if anchor is None:
        _failed.append(f"{tag}（找不到锚点）")
        return
    new_p = copy.deepcopy(anchor._p)
    anchor._p.addnext(new_p)
    p = Paragraph(new_p, anchor._parent)
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    run = p.add_run("▲ " + text)
    _set_font(run, size=10, bold=True, color=RED)
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    p.paragraph_format.line_spacing = 1.5
    _applied.append(tag)


def _tables(doc):
    """按文档流顺序返回表格列表。"""
    out = []
    for child in doc.element.body.iterchildren():
        if child.tag.endswith("}tbl"):
            out.append(Table(child, doc))
    return out


def _edit_cell(tbl, row_match, col, new_text, tag, *, col_match=0):
    """按某列内容定位行，改另一列的内容。"""
    for row in tbl.rows:
        if row.cells[col_match].text.strip() == row_match:
            cell = row.cells[col]
            para = cell.paragraphs[0]
            for r in list(para.runs):
                r._r.getparent().remove(r._r)
            run = para.add_run("▲ " + new_text)
            _set_font(run, size=9)
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            _applied.append(tag)
            return True
    _failed.append(f"{tag}（表里找不到行 {row_match}）")
    return False


def _insert_row_after(tbl, after_key, values, tag, *, col_match=0):
    """在 after_key 所在行之后插入一行（新增字段要插在字段表的合理位置）。"""
    import copy
    target = None
    for row in tbl.rows:
        if row.cells[col_match].text.strip().replace("▲ ", "") == after_key:
            target = row
            break
    if target is None:
        _failed.append(f"{tag}（找不到锚点行 {after_key}）")
        return False
    new_tr = copy.deepcopy(target._tr)
    target._tr.addnext(new_tr)
    from docx.table import _Row
    row = _Row(new_tr, tbl)
    for i, v in enumerate(values[: len(row.cells)]):
        para = row.cells[i].paragraphs[0]
        for r in list(para.runs):
            r._r.getparent().remove(r._r)
        run = para.add_run(("▲ " if i == 0 else "") + v)
        _set_font(run, size=9)
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    _applied.append(tag)
    return True


def _add_row(tbl, values, tag):
    row = tbl.add_row().cells
    for i, v in enumerate(values[: len(row)]):
        para = row[i].paragraphs[0]
        for r in list(para.runs):
            r._r.getparent().remove(r._r)
        run = para.add_run(("▲ " if i == 0 else "") + v)
        _set_font(run, size=9)
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    _applied.append(tag)


# ── 逐条应用 ────────────────────────────────────────────────────────────────

def main() -> None:
    doc = docx.Document(str(SRC))
    tbls = _tables(doc)
    # 前 3 张是新插的变更摘要表，原文档的表从第 4 张开始
    off = 3
    T = lambda i: tbls[off + i]  # noqa: E731  原文档表格编号

    # 【A】3.2 record_type 枚举补 3 个
    _edit_para(
        doc, "record_type 枚举：outpatient",
        "op_record 手术记录。",
        "senior_round 上级查房 / pre_op_summary 术前小结 / "
        "op_record 手术记录 / post_op_record 术后病程。",
        "A-record_type枚举",
    )

    # 【A】3.1 字段表新增 patient_no —— 必须真加进字段表，
    # 光写在摘要里不够：厂商是照字段表写代码的
    # 必填列写「否」而不是「是」：我方不传也照收（不会 40004 拒收），
    # 与 id_card 用同一套话术。若标「是」，摘要那边写的又是「建议必填」，
    # 两个口径对厂商是两种校验行为，必然回头来问。
    _insert_row_after(T(6), "hospital_code", [
        "patient_no", "string", "否（强烈建议）",
        "患者编号 / 病案号，同一患者跨就诊保持稳定。我方据此把同一个人的历次"
        "就诊归到同一份档案下——不传不会被拒收，但身份证也缺时该患者每次就诊"
        "都会新建档案，医生看不到其既往病历与过敏史。贵方取病案号 / 患者主索引号 / "
        "居民健康档案号均可，告知取的是哪个即可。我方仅用于档案查重，不回写",
        "BRID / 病案号对应项（贵方定，告知即可）",
    ], "A-patient_no字段")

    # 【A】3.3 target 说明（表 T10 第 3 列=说明）
    # 措辞注意：不能写「无 _record 后缀」——枚举里 course_record / op_record 等
    # 五个本来就以 _record 结尾，那样写会被理解成「要去掉后缀再填」，
    # 反而制造本条想根治的那个故障。改为「原样透传」+ 正例。
    _edit_cell(T(10), "target", 3,
               "刷新目标：把本次写入的 record_type 原样透传，不要做任何加/减"
               "后缀的加工。例如 record_type=outpatient 时 target 就填 "
               "outpatient；record_type=course_record 时就填 course_record",
               "A-target说明")

    # 【A】附录 A.3 示例去掉 _record 后缀
    _edit_para(doc, '"target": "outpatient_record"',
               '"target": "outpatient_record"', '"target": "outpatient"',
               "A-A3示例")

    # 【A】2.2 请求头表加 X-Request-Id
    _add_row(T(2), ["X-Request-Id",
                    "请求级幂等键。同一份文书的所有重发（含超时重试与我方对账补投）"
                    "该值恒定，建议贵方据此做一层请求去重；不参与签名计算"],
             "A-X-Request-Id")

    # 【A】7.2 握手失败 401 → 403
    _edit_para(doc, "连接地址：wss://mediscribe.cn",
               "鉴权失败我方拒绝握手（HTTP 401）",
               "鉴权失败我方拒绝 WebSocket 升级，贵方将收到 HTTP 403"
               "（不是 401，也没有 WS close code）",
               "A-握手403")

    # 【A】7.2 回写路由：补充换线说明
    _append_para_after(
        doc, "连接地址：wss://mediscribe.cn",
        "回写路由补充：常态下回写/刷新路由回接诊来源诊室；若该连接已断开或失效，"
        "我方会改投任意一条在线连接（msg_id 不变）。因此贵方任一客户端收到 "
        "record_writeback 都请无条件写库（写的是全院共享库，与诊室无关），"
        "不要按「是不是本机接诊的」过滤，否则该份病历会丢失。",
        "A-回写换线",
    )

    # 【A】7.4 补发与去重：拆上行/下行
    _rewrite_para(
        doc, "补发与去重：发送方对未收到 ack 的消息",
        "补发与去重（分方向）：① 上行（贵方→我方）：超时重发请保持 msg_id 不变，"
        "我方按 msg_id 幂等去重。② 下行（我方→贵方）：同一条消息的超时重发与"
        "换线重投 msg_id 不变，但连接断开后由我方对账任务发起的重投会使用新的 "
        "msg_id，故贵方判重必须以 2.5 的业务幂等键"
        "（visit_id + record_type + record_no）为准，不能只靠 msg_id。",
        "A-补发去重",
    )

    # 【B】第 4 章：删掉要字典（表 T11）
    _edit_cell(T(11), "需贵方提供", 1,
               "无需提供任何字典（与附录 B 已确认一致）。诊断以名称为准落库；"
               "科室以接诊推送随行的 dept_code + dept_name 成对下发为准。",
               "B-删字典要求")

    # 【B】3.1 agent_device_ip / mac（表 T6 说明列=3）
    _edit_cell(T(6), "agent_device_ip", 3,
               "选填，仅作排障留痕。方案 B 下回写按 WebSocket 连接自动路由回来源诊室，"
               "不依赖本字段，可不填", "B-agent_ip")
    _edit_cell(T(6), "agent_device_mac", 3,
               "保留字段，当前不使用，可不填", "B-agent_mac")

    # 【B】5.2 删极薄 Agent 段
    _rewrite_para(
        doc, "极薄 Agent 仅做两件事",
        "（本节原「诊室部署极薄 Agent 做内网中转」方案已取消）方案 B 下我方不调用"
        "贵方任何接口、不在院内部署任何程序，贵方无需预留内网部署位；"
        "仅当回落方案 A 时才需贵方提供公网可达的回写地址。",
        "B-删Agent",
    )

    # 【B】流程图第 ② 步
    _edit_para(
        doc, "①医生在 HIS 接诊病人",
        "设备IP", "医生工号", "B-流程图要素",
    )
    _append_para_after(
        doc, "关联主键：就诊流水号（visit_id）",
        "接入方式更正：原「诊室极薄 Agent 拉起嵌入工作台」的接入范式已下线。"
        "现为：我方收到推送后自动建档并按 doctor_code 派入对应医生的工作台；"
        "医生在已登录的浏览器工作台收到叫号提示后点击进入该患者。"
        "双方均无需安装任何客户端程序。",
        "B-接入方式",
    )

    # 【C】附录 A.2 示例补两个字段
    _edit_para(
        doc, '"record_type": "outpatient", "is_tcm"',
        '"record_type": "outpatient",',
        '"record_type": "outpatient", "record_no": "1", '
        '"full_text": "（病历完整文本）",',
        "C-A2示例补字段",
    )

    # 【C】3.2 status 口径（表 T8 说明列）
    _edit_cell(T(8), "status", 3,
               "写入状态。我方恒发 draft，即使该病历在 MediScribe 侧已由医生签发；"
               "贵方收到后落草稿/待确认态，医生可在 HIS 端继续编辑，"
               "点确认/保存后才真实落库。final 为后续预留，本期不会下发",
               "C-status口径")

    # 【C】3.1 dept_code / id_card 说明
    _edit_cell(T(6), "dept_code", 3,
               "本次接诊/挂号的科室编码，取贵方科室表「本地ID」列（如 1230024）。"
               "我方按此编码自动登记并作为本次接诊的科室（不是医生编制科室，"
               "医生跨科室坐班时以此为准），请确保同一科室取值稳定",
               "C-dept_code说明")
    # 与新增的 patient_no 行口径必须一致，否则同一张表里两行互相打架、
    # 厂商不知道档案查重到底靠哪个。这里把三级顺序一次讲清。
    _edit_cell(T(6), "id_card", 3,
               "身份证号。选填，但强烈建议推送：我方按「身份证 → patient_no → "
               "手机号+姓名」的顺序跨就诊认人，身份证优先级最高。"
               "三者全缺时同一位患者每次就诊都会新建档案，"
               "医生将看不到其既往病历与过敏史",
               "C-id_card说明")

    # 【C】3.1 幂等与体征补推
    _append_para_after(
        doc, "关于主诉：接诊是流程起点",
        "幂等与补推说明：幂等键为 hospital_code + visit_id（请保证同一次就诊两者完全"
        "一致）；重复推送复用已有接诊，若携带新的 doctor_code 会自动改派医生。"
        "体征与档案字段仅在该 visit_id 首次推送时生效，后续补推不再更新"
        "（避免覆盖医生已确认的内容），请在首次推送时尽量带齐。",
        "C-幂等与补推",
    )

    # 【C】3.1 宽容解析说明
    _append_para_after(
        doc, "体征与档案字段（上表 vitals.*",
        "取值容错：我方已对以下字段做宽容解析，贵方按 HIS 原生类型推送即可，"
        "不会因类型/编码差异导致整条推送被拒——性别接受 male/female/unknown、"
        "国标代码 1/2/9 与中文；初诊标志接受 bool 与 CZBZDM 代码 1/2；"
        "体征、身份证、手机号等接受 JSON 数字。无法识别的性别按 unknown 处理。",
        "C-宽容解析",
    )

    # 【C】2.2 重试规则限定方案
    _edit_para(
        doc, "重试规则：nonce 为一次性",
        "重试规则：",
        "重试规则（本条适用于方案 A）：",
        "C-重试规则限定",
    )
    # 锚点用 "nonce 为一次性" 而不是 "重试规则："——上一步刚把
    # "重试规则：" 改成了 "重试规则（本条适用于方案 A）："，原锚点已不存在
    _append_para_after(
        doc, "nonce 为一次性",
        "方案 B 补充：WebSocket 通道的业务消息不做 nonce 去重，防重放依赖"
        "握手鉴权 + ±5 分钟时间戳窗口 + msg_id 幂等；重发时 timestamp/nonce/sign "
        "重新生成而 msg_id 保持不变（见 7.4）。故 40006 不会出现在 WS 的 ack 中。",
        "C-方案B重试",
    )

    # 【C】7.4 心跳发起方
    _rewrite_para(
        doc, "心跳：任一方每 30 秒发送",
        "心跳：由我方每 30 秒下发 type=ping，贵方收到后立即回 type=pong 即可"
        "（贵方无需主动发起心跳；主动发我方也会回 pong）。"
        "任一方连续 90 秒未收到任何消息即判定连接失效。",
        "C-心跳发起方",
    )

    # 【C】第 6 章联调步骤
    _rewrite_para(
        doc, "2. 接诊推送连通：HIS 能 POST",
        "2. 医生名单交换：贵方提供参与联调的临床医生名单（姓名 + 全部工号，"
        "含本人门诊/住院/助理等多个工号），我方完成开户并绑定后方可推送；"
        "未注册的工号推送将返回 code=40007。",
        "C-联调医生名单",
    )
    _rewrite_para(
        doc, "3. 回写连通：MediScribe 能 POST",
        "3. 通道连通（方案 B）：贵方客户端 wss 握手鉴权通过；上行 encounter_admit "
        "我方回 ack code=0 且患者出现在医生工作台队列；下行 record_writeback / "
        "record_refresh 贵方落库并回 ack；ping/pong 心跳与断线自动重连正常。",
        "C-联调方案B",
    )

    # 【C】附录 A.1 dept_code 示例
    _edit_para(doc, '"dept_code": "0101"',
               '"dept_code": "0101", "dept_name": "全科(安元)"',
               '"dept_code": "1230024", "dept_name": "全科"',
               "C-A1示例科室")

    # 【C】2.5 空值语义
    _rewrite_para(
        doc, "字段语义：字段缺省（不传）= 不更新该项",
        "字段语义：字段缺省（不传）= 不更新该项；传空串（\"\"）= 清空该项。"
        "对象型字段（record / vitals / meta）与数组型字段（diagnoses）："
        "整体缺省或为空（{} / []）均表示不更新该组，不做清空。"
        "另：MediScribe 本期不使用空串语义——字段为空即不下发；"
        "如需清空 HIS 中某项，请在 HIS 侧人工编辑。",
        "C-空值语义",
    )

    # 【C】meta 字段说明（表 T8）
    _edit_cell(T(8), "meta.generated_at", 3,
               "本次回写报文的生成时间（北京时间；重发/补投时会刷新）",
               "C-generated_at")
    _edit_cell(T(8), "meta.doctor_code", 3,
               "该就诊的接诊医生工号（取自 3.1 接诊推送的 doctor_code）",
               "C-doctor_code")

    # 【C】2.1 时间格式补时区
    _edit_cell(T(1), "时间格式", 1,
               "YYYY-MM-DD HH:mm:ss；所有时间均为北京时间（UTC+8），不带时区后缀",
               "C-时区说明")

    # ══ 发文前最后一轮核对补漏（2026-08-15）══════════════════════════════
    # 以下几处是「同一件事在文档里有两种说法」或「示例与正文互相否定」。
    # 厂商照示例做是最常见的开发方式，示例错比正文错更致命。

    # 【B】3.2 诊断说明末尾还在要字典 —— 与第 4 章、附录 B 直接矛盾
    _edit_para(doc, "诊断说明：MediScribe 当前阶段以诊断名称",
               "（请贵方提供本院诊断字典，见附录 B）",
               "（诊断字典无需贵方提供，见第 4 章与附录 B）",
               "B-诊断字典残留")

    # 【B】5.3 还留着「极薄 Agent 仍保留拉起工作台职责」—— 与 1.3、5.2 矛盾，
    # 厂商读到会以为院内仍需部署我方程序，进而去申请部署位、排 IT 审批
    _edit_para(doc, "若采用 WebSocket 通道，5.1 与 5.2 的网络要求整体简化",
               "；我方亦无需在内网中转回写（极薄 Agent 仅保留「拉起嵌入工作台」职责，"
               "不再承担内网转发）。",
               "；我方亦无需在内网中转回写，院内不部署我方任何程序。",
               "B-5.3极薄Agent残留")

    # 【B】1.3 流程图第 ② 步没改过，却被整块黄色高亮盖住，看着像已改
    _edit_para(doc, "②MediScribe 据此识别当前病人",
               "②MediScribe 据此识别当前病人；诊室「极薄 Agent」拉起嵌入工作台",
               "②MediScribe 自动建档并按 doctor_code 派入该医生的浏览器工作台；"
               "医生点击叫号提示进入",
               "B-流程图②")

    # 【B】5.2 可达性表里「诊室部署极薄 Agent」那一行还在（表在上、否定在下，
    # 厂商扫表格时只会看到表格）
    for _t in tbls:
        for _row in _t.rows:
            if "极薄" in _row.cells[-1].text:
                _c = _row.cells[-1]
                _p = _c.paragraphs[0]
                for _r in list(_p.runs):
                    _r._r.getparent().remove(_r._r)
                _run = _p.add_run("▲ （本方案已取消，见本节下方说明）")
                _set_font(_run, size=9)
                _run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                _applied.append("B-5.2表格极薄Agent行")

    # 【A】附录 A.4 两条 ack 示例补齐签名字段。我方对 ack 强制验签，
    # 未签名的 ack 会被静默丢弃 → 回写重发 3 次后断连 → 整份病历回写失败
    for _old, _new in (
        ('{ "type": "ack", "msg_id": "ms-ack-0001",\n'
         '  "payload":',
         '{ "type": "ack", "msg_id": "ms-ack-0001",\n'
         '  "app_id": "mediscribe", "timestamp": "1770000000100", '
         '"nonce": "b2c3d4", "sign": "…",\n'
         '  "payload":'),
        ('{ "type": "ack", "msg_id": "his-ack-0007",\n'
         '  "payload":',
         '{ "type": "ack", "msg_id": "his-ack-0007",\n'
         '  "app_id": "his_wanda", "timestamp": "1770000001100", '
         '"nonce": "e5f6a7", "sign": "…",\n'
         '  "payload":'),
    ):
        _edit_para(doc, _old.split("\n")[0], _old, _new, "A-ack示例补签名")

    _append_para_after(
        doc, "【已确认】连接拓扑：WS 代码加在贵方全科客户端内",
        "补充：来源诊室连接失效时，我方会改投任意一条在线连接，"
        "故贵方任一客户端收到 record_writeback 均须无条件写库，"
        "不要按「是不是本机接诊的」过滤（详见 7.2）。",
        "A-附录B无条件写库",
    )

    # 【A】附录 A.1 示例：补 patient_no / id_card，机构码与参考实现对齐，
    # 删掉已改为「可不填」的 agent_device_ip（示例带着等于否定正文）
    _edit_para(doc, '"visit_id": "20260514001401"',
               '"hospital_code": "H32050701121",',
               '"hospital_code": "H33052300001",\n'
               '  "patient_no": "BA20260514001",',
               "A-A1示例补patient_no")
    _edit_para(doc, '"visit_id": "20260514001401"',
               '"patient_name": "测试111", "gender": "male", '
               '"birth_date": "2000-01-01",',
               '"patient_name": "测试111", "gender": "male", '
               '"birth_date": "2000-01-01",\n'
               '  "id_card": "330523200001010038", "phone": "13800138000",',
               "A-A1示例补id_card")
    _edit_para(doc, '"visit_id": "20260514001401"',
               '\n  "agent_device_ip": "172.16.2.58",', "",
               "B-A1示例删device_ip")

    # 【A】7.2 握手 nonce 的一次性语义此前完全没写。把握手 URL 算好存成配置
    # 或成员变量、断线后原样重连，是很自然的写法，结果是每次重连都被 403，
    # 而 403 与 appId 错、签名错的表现完全一样，现场无从区分
    _append_para_after(
        doc, "连接地址：wss://mediscribe.cn/api/v1/his/ws",
        "握手 nonce 为一次性，我方在时间窗内做去重：每次（重）连都必须重新生成 "
        "timestamp / nonce / sign，不得把算好的握手 URL 缓存下来原样复用，"
        "否则同样返回 HTTP 403，表现为「首次能连、断线后再也连不上」。"
        "（随附的两份参考实现在 connect() 内每次现算，可直接对照。）",
        "A-握手nonce一次性",
    )

    # 【C】第 6 章收尾还在要贵方提供「测试凭证」，但方案 B 凭证是单向的
    _edit_para(doc, "请贵方提供：测试环境 HIS + 测试凭证",
               "请贵方提供：测试环境 HIS + 测试凭证，先在测试环境跑通再上生产。",
               "请贵方提供：测试环境的 HIS 客户端与测试用就诊数据；"
               "测试凭证由我方分配（方案 B 为单向一套，贵方无需为我方分配）。"
               "先在测试环境跑通再上生产。",
               "C-第6章测试凭证")

    # 【C】1.2 住院文书括号列举还是老四类，与补齐后的 record_type 枚举对不上
    _edit_para(doc, "本回写接口为统一入口",
               "住院各类文书（入院记录/病程/出院/手术）",
               "住院各类文书（见 3.2 record_type 枚举）",
               "C-1.2住院文书列举")

    # 【C】扉页修订说明开头仍是「v1.2 修订：」，版本号却已是 v1.3
    _edit_para(doc, "v1.2 修订：通道定稿",
               "v1.2 修订：通道定稿",
               "v1.3 修订：共 %d 处，逐条见文首《本次修订说明》。"
               "（v1.2 修订：通道定稿" % len(CHANGES),
               "C-扉页版本说明")

    # 结尾版本号。注意不打 ▲：全文 ▲ 的含义是「本次规范条目修订」，
    # 版本号变更不是 26 条中的任何一条，标了会让按 ▲ 数条目的人对不上数
    _p_tail = _find_para(doc, "— 本文档 v1.2，供双方评审")
    if _p_tail is not None:
        for _r in _p_tail.runs:
            _r.text = _r.text.replace("本文档 v1.2", "本文档 v1.3")
        _applied.append("尾注版本号（不打▲）")
    else:
        _failed.append("尾注版本号")

    doc.save(str(SRC))

    print(f"已应用 {len(_applied)} 项：")
    for t in _applied:
        print("   ✓", t)
    if _failed:
        print(f"\n⚠ 未命中 {len(_failed)} 项（需人工确认）：")
        for t in _failed:
            print("   ✗", t)
    sys.exit(1 if _failed else 0)


if __name__ == "__main__":
    main()
