# -*- coding: utf-8 -*-
"""生成《MediScribe HIS 接口规范 v1.5 增补说明》docx（阶段6，2026-08-21）。

产出：D:/文件/项目文档/MediScribe/MediScribe_HIS接口规范_v1.5_增补说明.docx
交付对象：HIS 厂商（随联调发出）。

铁律（记忆 feedback_spec_examples_over_prose）：示例报文必须与
his_adapter/writeback_builder.py 的真实输出逐字段一致——本脚本的示例块
直接按 _build_diagnoses 的产出结构手工核对（字段名/类型/枚举与代码同源），
修改 builder 后必须重新生成本文档。
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt

OUT = Path(r"D:/文件/项目文档/MediScribe/MediScribe_HIS接口规范_v1.5_增补说明.docx")

EXAMPLE = """{
  "visit_id": "MZ20260821001",
  "record_type": "discharge_record",
  "record_no": "1",
  "is_tcm": true,
  "status": "draft",
  "record": { "...（与 v1.4 一致，略）" : "" },
  "vitals":  { "...（与 v1.4 一致，略）" : "" },
  "diagnoses": [
    {
      "name": "原发性高血压",
      "is_primary": true,
      "category": "western",
      "admission_condition": "有",
      "code": "I10.x09",
      "code_type": "ICD10"
    },
    {
      "name": "2型糖尿病",
      "is_primary": false,
      "category": "western",
      "admission_condition": "临床未确定",
      "code": "E11.900",
      "code_type": "ICD10"
    },
    {
      "name": "眩晕",
      "is_primary": false,
      "category": "tcm_disease",
      "admission_condition": "有",
      "code": "A17.07",
      "code_type": "GB95"
    },
    {
      "name": "肝阳上亢证",
      "is_primary": false,
      "category": "tcm_syndrome",
      "code": "B04.02.01.04.02.01",
      "code_type": "GB95"
    }
  ],
  "meta": { "...（与 v1.4 一致，略）" : "" }
}"""


def main() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(10.5)

    doc.add_heading("MediScribe HIS 接口规范 v1.5 增补说明", level=0)
    doc.add_paragraph("——诊断编码能力上线（对 v1.4 的向后兼容增补）　2026-08-21")

    doc.add_heading("一、背景与兼容性声明", level=1)
    doc.add_paragraph(
        "v1.4 规范表 11 曾注明：诊断以名称（name）为准、不下发 code/code_type，"
        "编码能力后续迭代补充。本增补即该承诺的兑现：我方已完成诊断结构化与"
        "编码字典接入，病历回写接口（3.2）的 diagnoses[] 数组开始携带编码。"
    )
    p = doc.add_paragraph()
    r = p.add_run(
        "依据 v1.4 §2.5 向后兼容条款（新增可选字段/枚举值不构成破坏性变更，"
        "贵方忽略未识别字段即可），本次增补【无需贵方任何开发改造】。"
        "如贵方希望利用编码直接落库/对接病案首页，可按下述字段说明采用。"
    )
    r.bold = True

    doc.add_heading("二、变更明细", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "字段"
    hdr[1].text = "变化"
    hdr[2].text = "说明"
    rows = [
        ("diagnoses[].code", "开始有值下发（可选）",
         "诊断编码；与 name 成对。我方系统内医生经字典联想选定或名称精确匹配自动补码；"
         "无法确定编码时该字段缺省（宁缺勿错）"),
        ("diagnoses[].code_type", "开始有值下发（可选）",
         "ICD10 = 疾病诊断编码（国家医保版 2.0 口径）；"
         "GB95 = 中医病证分类与代码（枚举字面量沿用 v1.4；数据实际依据已升级的"
         "GB/T 15657-2021 医保版，含疾病名 A 码与证候名 B 码——如贵方需要区分"
         "或采用其他标识，请联调时提出）"),
        ("diagnoses[].is_primary", "语义升级",
         "此前=第一个非空诊断；现在=医生在工作台显式勾选的主要诊断"
         "（按'本次医疗资源消耗最多、对健康危害最大'原则），全数组恰一条为 true"),
        ("diagnoses[]（western）", "可多条",
         "西医诊断从单条文本升级为逐条条目（主要诊断+其他诊断/合并症），"
         "顺序即主诊断在前的展示序"),
        ("diagnoses[].admission_condition", "新增可选字段",
         "入院病情（病案首页逐条诊断标志位）：有 / 临床未确定 / 情况不明 / 无；"
         "仅住院场景且医生已标注时下发，门急诊缺省"),
    ]
    for a, b, c in rows:
        cells = table.add_row().cells
        cells[0].text = a
        cells[1].text = b
        cells[2].text = c

    doc.add_heading("三、示例报文（节选 diagnoses 段）", level=1)
    doc.add_paragraph(
        "以下示例与我方回写实现逐字段一致（字段缺省语义仍按 §2.5：缺省=不更新）："
    )
    code_p = doc.add_paragraph(EXAMPLE)
    for run in code_p.runs:
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)

    doc.add_heading("四、贵方侧建议（可选，非必须）", level=1)
    doc.add_paragraph(
        "1. 如贵方病案首页模块支持按编码落库，可直接采用 code/code_type，"
        "免去名称匹配本院诊断库的环节；\n"
        "2. admission_condition 可直接映射贵方首页'入院病情'栏；\n"
        "3. 以上均为增量利用，不采用亦不影响现有联调口径与流程。"
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"已生成 {OUT}")


if __name__ == "__main__":
    main()
